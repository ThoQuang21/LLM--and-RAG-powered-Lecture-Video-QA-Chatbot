import os
import sys
import time
import random
import datetime
from typing import Dict, Optional, Any

# --- CÁC THƯ VIỆN CẦN CÀI ĐẶT ---
# 1. Hãy chạy ô này trước:
# !pip install openai-whisper python-docx yt-dlp

try:
    import whisper # Đây là whisper "cơ bản"
    from docx import Document
except ImportError:
    print("Vui lòng chạy ô lệnh sau để cài đặt:")
    print("!pip install openai-whisper python-docx yt-dlp")
    sys.exit(1)

# ==============================================================================
# PHẦN 1: CÁC HÀM HỖ TRỢ VÀ LOGIC CHÍNH
# ==============================================================================

def format_timestamp(seconds: float) -> str:
    """
    Chức năng: Chuyển đổi giây (float) sang định dạng [MM:SS,ms] hoặc [HH:MM:SS,ms].
    (Hàm này giữ nguyên)
    """
    total_seconds = int(seconds)
    milliseconds = int((seconds - total_seconds) * 1000)

    hours, remainder = divmod(total_seconds, 3600)
    minutes, seconds_val = divmod(remainder, 60)

    if hours > 0:
        return f"[{hours:02}:{minutes:02}:{seconds_val:02},{milliseconds:03}]"
    else:
        return f"[{minutes:02}:{seconds_val:02},{milliseconds:03}]"

def load_whisper_model(model_name: str = "base") -> Optional[whisper.Whisper]:
    """
    Chức năng: Tải model Whisper CƠ BẢN (openai-whisper).
    Bạn có thể đổi 'base' thành 'small', 'medium' để chính xác hơn (nhưng chậm hơn).
    """
    print(f"\n--- ĐANG TẢI MODEL WHISPER CƠ BẢN ({model_name}) ---")

    try:
        # Tự động dùng GPU (cuda) nếu có
        model = whisper.load_model(model_name)
        print("Tải model Whisper cơ bản thành công.")
        return model
    except Exception as e:
        print(f"[LỖI NGHIÊM TRỌNG] Không thể tải model Whisper: {e}")
        return None

def download_audio(video_url: str, output_path: str, cookies_path: str) -> bool:
    """
    Chức năng: Tải audio từ video_url về output_path (dạng mp3) bằng yt-dlp.
    (Hàm này giữ nguyên)
    """
    print(f"   Đang tải audio từ: {video_url.split('&t=')[0]}")
    url_to_download = video_url.split('&t=')[0]

    cookies_cmd = f'--cookies "{cookies_path}"' if cookies_path else ""
    command = (
        f'yt-dlp --no-check-certificate {cookies_cmd} '
        f'-x --audio-format mp3 '
        f'-o "{output_path}" '
        f'"{url_to_download}"'
        " > /dev/null 2>&1" # Ẩn output
    )

    return_code = os.system(command)

    if return_code != 0:
         print(f"   [LỖI] yt-dlp thất bại (return code: {return_code}).")
         return False
    if not os.path.exists(output_path):
        print("   [LỖI] File audio không được tạo sau khi tải.")
        return False

    print("   Tải audio thành công.")
    return True

def transcribe_audio(model: whisper.Whisper, audio_path: str) -> Optional[Dict[str, Any]]:
    """
    Chức năng: Chạy model.transcribe() (của whisper cơ bản).
    """
    print("   Đang chạy AI Speech-to-Text (Whisper cơ bản)...")
    try:
        # Ép nhận diện tiếng Việt
        result = model.transcribe(audio_path, language="vi")
        return result
    except Exception as e:
        print(f"   [LỖI] Xảy ra lỗi trong quá trình transcribe: {e}")
        return None

def save_transcript_to_docx(result: Dict[str, Any], save_path: str):
    """
    Chức năng: Lưu kết quả từ "whisper cơ bản" ra file .docx.
    Định dạng: [START_TIME] -> [END_TIME] Text
    """
    print(f"   Đang lưu vào file: {save_path}")
    doc = Document()

    # "whisper cơ bản" dùng key là "segments"
    segments = result.get("segments", [])

    if not segments:
        print("   [CẢNH BÁO] Không tìm thấy 'segments'. Đang lưu text đầy đủ.")
        doc.add_paragraph(result.get("text", "Không có nội dung."))
    else:
        print(f"   Đã tìm thấy {len(segments)} đoạn transcript. Đang xử lý...")
        for segment in segments:
            # "whisper cơ bản" lưu start/end trực tiếp
            start_time = segment.get('start')
            end_time = segment.get('end')
            text = segment.get('text', '').strip()

            # --- SỬA LỖI 'NoneType' ---
            # (Giữ lại logic sửa lỗi của ViWhisper để phòng hờ)
            if not text or start_time is None:
                continue
            if end_time is None:
                print(f"   [CẢNH BÁO] Segment cuối không có end_time. Dùng tạm start_time.")
                end_time = start_time
            # --- KẾT THÚC SỬA LỖI ---

            # Định dạng lại thời gian
            formatted_start = format_timestamp(start_time)
            formatted_end = format_timestamp(end_time)

            transcript_line = f"{formatted_start} -> {formatted_end} {text}"
            doc.add_paragraph(transcript_line)

    doc.save(save_path)
    file_name = os.path.basename(save_path)
    print(f"   [THÀNH CÔNG]: Đã lưu {file_name}")

def process_videos(video_map: Dict, save_dir: str, model: whisper.Whisper, temp_audio_file: str, cookies_path: str):
    """
    Chức năng: Vòng lặp chính xử lý từng video trong video_map.
    (Hàm này giữ nguyên)
    """
    print("\n--- BẮT ĐẦU XỬ LÝ VIDEO ---")

    for file_name_docx, video_url in video_map.items():
        print(f"\n[ĐANG XỬ LÝ]: {file_name_docx}")
        full_save_path = os.path.join(save_dir, file_name_docx)

        try:
            # 1. Tải Audio
            if not download_audio(video_url, temp_audio_file, cookies_path):
                print(f"   [BỎ QUA] Lỗi tải audio cho {file_name_docx}.")
                continue

            # 2. Chuyển đổi (Transcribe)
            result = transcribe_audio(model, temp_audio_file)
            if result is None:
                print(f"   [BỎ QUA] Lỗi transcribe cho {file_name_docx}.")
                continue

            # 3. Lưu file
            save_transcript_to_docx(result, full_save_path)

            # 4. Nghỉ ngơi
            sleep_time = random.uniform(5, 15) # Giảm thời gian nghỉ
            print(f"🕓 Nghỉ {sleep_time:.1f} giây...")
            time.sleep(sleep_time)

        except Exception as e:
            print(f"   [LỖI TỔNG QUÁT] Xảy ra lỗi với file {file_name_docx}: {e}")

        finally:
            # 5. Xóa file tạm
            if os.path.exists(temp_audio_file):
                os.remove(temp_audio_file)

    print("\n--- HOÀN TẤT TẤT CẢ VIDEO ---")

# ==============================================================================
# PHẦN 2: HÀM MAIN ĐỂ CHẠY
# ==============================================================================

def main():
    """
    Hàm chính điều phối toàn bộ
    """
    # 1. THIẾT LẬP CẤU HÌNH (Tùy chỉnh ở đây)
    print("--- 1. ĐANG KHỞI TẠO BIẾN ---")

    VIDEO_MAP = {
    "video_C3_3_1_transcript.docx": "https://www.youtube.com/watch?v=KeNRQw9j_ps",
}

    DRIVE_SAVE_PATH = "/content/drive/MyDrive/DL_RAG_Video_main/Transcripts"
    TEMP_AUDIO_FILE = "/content/temp_audio.mp3"
    COOKIES_PATH = "" # Để "" nếu không dùng cookies

    os.makedirs(DRIVE_SAVE_PATH, exist_ok=True)
    print(f"Sẵn sàng lưu file vào: {DRIVE_SAVE_PATH}")

    # 2. TẢI MODEL
    # Chọn model: "tiny", "base", "small", "medium", "large"
    # "base" hoặc "small" là đủ dùng.
    asr_model = load_whisper_model("medium")

    if asr_model is None:
        print("Không thể tải model. Thoát chương trình.")
        sys.exit(1)

    # 3. CHẠY XỬ LÝ
    process_videos(
        video_map=VIDEO_MAP,
        save_dir=DRIVE_SAVE_PATH,
        model=asr_model,
        temp_audio_file=TEMP_AUDIO_FILE,
        cookies_path=COOKIES_PATH
    )
