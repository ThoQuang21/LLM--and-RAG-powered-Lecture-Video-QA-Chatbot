import google.generativeai as genai
from google.colab import userdata
from docx import Document
import os
import time
from typing import Set, Dict, Optional

# ==============================================================================
# PHẦN 1: CÁC HÀM LOGIC 
# ==============================================================================

def initialize_gemini(api_key: str) -> (genai.GenerativeModel | None):
    """
    Chức năng: Kết nối đến Google Gemini.
    """
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-2.5-flash')
        print("Đã kết nối Gemini thành công (Model: 2.5-Flash).")
        return model
    except Exception as e:
        print(f"Lỗi: Không thể kết nối Gemini: {e}")
        return None

def read_entire_docx_file(source_path: str) -> (str | None):
    """
    Chức năng: Đọc TOÀN BỘ file .docx và gộp tất cả text lại thành 1 chuỗi.
    """
    try:
        doc = Document(source_path)
        original_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

        if not original_text:
            print(" <em> File gốc rỗng, có thể đang được ghi.</em>")
            return None
        return original_text
    except Exception as e:
        print(f" <em> [LỖI ĐỌC FILE] {source_path}: {e}</em>")
        return None

def call_gemini_on_full_text(
    model: genai.GenerativeModel,
    full_text: str,
    prompt_template: str
) -> (str | None):
    """
    Chức năng: Gửi TOÀN BỘ text cho Gemini 1 LẦN DUY NHẤT.
    """
    try:
        print(" <em> Đang gửi cho Gemini sửa (1 lần gọi)...</em>")
        prompt_with_text = prompt_template.format(TEXT_TO_CLEAN=full_text)
        response = model.generate_content(prompt_with_text)
        return response.text.strip()
    except Exception as e:
        print(f" <em> [LỖI GEMINI] Xảy ra lỗi khi gọi API: {e}</em>")
        return None

def save_cleaned_text(cleaned_text: str, cleaned_path: str):
    """
    Chức năng: Lưu 1 khối text duy nhất vào file .docx.
    """
    try:
        print(f" <em> Đang LƯU file đã sửa vào: {cleaned_path}</em>")
        new_doc = Document()
        new_doc.add_paragraph(cleaned_text)
        new_doc.save(cleaned_path)
    except Exception as e:
        print(f" <em> [LỖI LƯU FILE] Không thể lưu {cleaned_path}: {e}</em>")

def start_watcher(
    files_to_clean_set: Set[str],
    source_dir: str,
    cleaned_dir: str,
    model: genai.GenerativeModel,
    prompt_template: str
):
    """
    Chức năng: Vòng lặp "canh gác" (watcher) chính.
    """
    print(f"\n--- Bắt đầu chế độ 'chờ' (watcher) ---")
    print(f"Đang theo dõi {len(files_to_clean_set)} file trong thư mục: {source_dir}")
    print(f"Kết quả sẽ được lưu vào: {cleaned_dir}")

    try:
        while True:
            if not files_to_clean_set:
                print("\nĐã xử lý tất cả file. Tắt script.")
                break

            files_to_check = list(files_to_clean_set)
            found_a_file = False

            for file_name in files_to_check:
                source_path = os.path.join(source_dir, file_name)
                cleaned_path = os.path.join(cleaned_dir, file_name)

                # 4.1. KIỂM TRA XEM ĐÃ SỬA CHƯA
                if os.path.exists(cleaned_path):
                    print(f"[ĐÃ XONG] Bỏ qua {file_name}, file đã được dọn dẹp từ trước.")
                    files_to_clean_set.remove(file_name)
                    found_a_file = True
                    continue

                # 4.2. KIỂM TRA FILE GỐC ĐÃ XUẤT HIỆN CHƯA
                if os.path.exists(source_path):
                    print(f"\n[PHÁT HIỆN FILE GỐC]: {file_name}")
                    found_a_file = True

                    try:
                        time.sleep(5) # Chờ file ghi

                        # 4.3. ĐỌC FILE GỐC 
                        original_text = read_entire_docx_file(source_path)
                        if not original_text:
                            continue

                        # 4.4. GỌI GEMINI SỬA 
                        cleaned_text = call_gemini_on_full_text(model, original_text, prompt_template)
                        if not cleaned_text:
                            print(" <em> Gemini lỗi, sẽ thử lại ở vòng lặp sau.</em>")
                            time.sleep(10) # Nghỉ 10s nếu Gemini lỗi
                            continue

                        # 4.5. LƯU VÀO THƯ MỤC MỚI 
                        save_cleaned_text(cleaned_text, cleaned_path)
                        print(f" <em> [THÀNH CÔNG]: Đã tạo file đã sửa cho {file_name}</em>")

                        # 4.6. XÓA KHỎI DANH SÁCH "CẦN LÀM"
                        files_to_clean_set.remove(file_name)

                        # Thêm thời gian nghỉ để tránh lỗi 503 (quá tải API)
                        print("🕓 Nghỉ 15 giây để tránh rate limit của Gemini...")
                        time.sleep(15)

                    except Exception as e:
                        print(f" <em> [LỖI] Xảy ra lỗi khi xử lý {file_name}: {e}</em>")
                        print(" <em> Sẽ thử lại ở vòng lặp sau.</em>")

            # 4.7. NGHỈ (NẾU KHÔNG TÌM THẤY GÌ)
            if not found_a_file and files_to_clean_set:
                sleep_time = 60
                print(f"... (Còn {len(files_to_clean_set)} file chưa xuất hiện) ... Đang chờ {sleep_time} giây trước khi quét lại ...")
                time.sleep(sleep_time)
            elif found_a_file:
                time.sleep(1) # Quét lại ngay

    except KeyboardInterrupt:
        print("\nĐã dừng script.")

# ==============================================================================
# PHẦN 3: HÀM MAIN ĐỂ CHẠY
# ==============================================================================

def main():
    """
    Hàm chính điều phối toàn bộ
    """

    # 1. API KEY
    GOOGLE_API_KEY = 'AIzaSyBD27hwT7Zu1yACDlbR1sEoVDKww2T2Cuo'

    # 2. ĐỊNH NGHĨA PROMPT (Prompt mới của bạn)
    PROOFREAD_PROMPT = """
Bạn là một trợ lý biên tập viên tiếng Việt xuất sắc.
Nhiệm vụ của bạn là nhận một đoạn văn bản (transcript của bài giảng) và thực hiện các việc sau:

1.  **Sửa lỗi chính tả:** Sửa tất cả các lỗi gõ sai, sai từ.
2.  **Ngữ pháp và dấu câu:** Thêm dấu câu (dấu chấm, phẩy, chấm hỏi) một cách hợp lý để câu văn dễ đọc và đúng ngữ pháp, tự động ngắt đoạn nếu hợp lí hay chuyển ý.
3.  **KHÔNG THAY ĐỔI Ý NGHĨA:** Tuyệt đối không được thêm nội dung mới, không bình luận, không tóm tắt, không thay đổi ý nghĩa gốc của câu.
4.  **Giữ nguyên từ chuyên ngành:** Nếu có từ chuyên ngành (ví dụ: RAG, AI, VectorDB), hãy giữ nguyên chúng.
5.  **sau khi ngắt đoạn thì thời timestampe cho các đoạn dựa trên các câu đã gộp và tôi chỉ muốn có thời gian bắt đầu của các đoạn
6.  **BẮT BUỘC Văn bản trả về phải được gộp lại thành các đoạn văn gồm vài câu cùng nghĩa, không được gộp luôn tuồn tất cả câu thành 1 đoạn duy nhất
Hãy trả về CHỈ văn bản đã được sửa sạch đẹp.

**Văn bản gốc cần sửa:**
---
{TEXT_TO_CLEAN}
---

**Văn bản đã sửa:**
"""

    # 3. ĐỊNH NGHĨA ĐƯỜNG DẪN
    SOURCE_DIR = "/content/drive/My Drive/DL_RAG_Video_main/Transcripts"
    CLEANED_DIR = "/content/drive/My Drive/DL_RAG_Video_main/Transcripts_Cleaned/"
    os.makedirs(CLEANED_DIR, exist_ok=True)

    # !!! QUAN TRỌNG: Bạn cần định nghĩa VIDEO_MAP ở đây
    # (Lấy từ Script 1)
    VIDEO_MAP: Dict[str, str] = {
        "video_C1_1_transcript.docx": "URL1",
        "video_C1_2_transcript.docx": "URL2",
        "video_C2_1_transcript.docx": "URL3",
        "video_C2_2_1_transcript.docx": "URL4",
        # Thêm các file .docx và URL video vào đây
    }

    files_to_clean_set = set(VIDEO_MAP.keys())

    # 4. KHỞI TẠO GEMINI
    llm_model = initialize_gemini(GOOGLE_API_KEY)

    if not llm_model:
        print("Không thể khởi tạo Gemini. Thoát chương trình.")
        return

    # 5. CHẠY VÒNG LẶP XỬ LÝ
    start_watcher(
        files_to_clean_set=files_to_clean_set,
        source_dir=SOURCE_DIR,
        cleaned_dir=CLEANED_DIR,
        model=llm_model,
        prompt_template=PROOFREAD_PROMPT
    )

# Chạy hàm main khi script được thực thi
if __name__ == "__main__":
    main()
